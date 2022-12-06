import os
import sys
import docx

def get_list_of_word_files(path):
    assert(os.path.exists(path))
    assert(os.path.isdir(path))
    tmp = [i for i in os.listdir(path) if (os.path.isfile(i) and i.split('.')[-1] == 'docx')]
    assert(tmp)
    return tmp



def change_style(dir, example):
    arr = get_list_of_word_files(dir)
    assert(example in arr)
    s = docx.Document(os.path.join(dir, example)).paragraphs[0].style

    for doc_file in arr:
        if doc_file != example:
            doc = docx.Document(os.path.join(dir, example))
            for paragraph in doc.paragraphs:
                paragraph.style = s
            doc.save(os.path.join(dir, doc_file))





if __name__ == '__main__':
    if len(sys.argv) == 1:
        basefilename = input("template.docx file name: ")
        folder = input("folder name")
    elif len(sys.argv) == 2:
        basefilename = sys.argv(1)
        folder = "."

    else:
        basefilename = sys.argv(1)
        folder = sys.argv[2]

    change_style(folder, basefilename)




